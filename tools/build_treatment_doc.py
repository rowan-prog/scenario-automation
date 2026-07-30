# -*- coding: utf-8 -*-
"""회차 트리트먼트 docx 빌더 (phase_p 부가 산출 · 2026-07-29 신설)

기획안 양식과 별개로 작가 전달용 트리트먼트를 뽑는다.
무료회차 = 기획안 spec의 `#FILL 40` 블록을 그대로 인용(상세)
유료회차 = `_paid_run_*.txt` (한 줄 = `NN화|본문 2~3문장`)

사용: python tools/build_treatment_doc.py <spec.txt> <paid_run.txt> <out.docx> "<제목>" "<부제>"
"""
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def free_run(path):
    """기획안 spec(#FILL 40 블록) 또는 별도 free-run 텍스트 파일 둘 다 받는다."""
    s = open(path, encoding='utf-8').read()
    if '#FILL 40' in s:
        s = s[s.index('#FILL 40\n') + len('#FILL 40\n'):s.index('#FILL 46')]
    return [l for l in s.split('\n') if l.strip()]


def build(spec_path, paid_path, out_path, title, subtitle):
    doc = Document()

    def para(text, bold=False, size=10, space=6):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        p.paragraph_format.space_after = Pt(space)

    para(title, bold=True, size=15, space=2)
    para(subtitle, size=9, space=14)

    para('무료회차 (1~8화) — 상세', bold=True, size=12, space=8)
    for line in free_run(spec_path):
        if re.match(r'^\d+화$', line.strip()):
            para(line, bold=True, size=11, space=8)
        elif line.startswith('클리프행어:'):
            para(line, size=10, space=14)
        else:
            para(line, size=10, space=2)

    para('유료회차 (9~50화) — 회차별 요약', bold=True, size=12, space=8)
    for row in open(paid_path, encoding='utf-8').read().strip().split('\n'):
        ep, body = row.split('|', 1)
        para(ep, bold=True, size=11, space=2)
        para(body, size=10, space=10)

    doc.save(out_path)
    print(f'OK: {out_path}')


if __name__ == '__main__':
    if len(sys.argv) != 6:
        raise SystemExit('사용: python tools/build_treatment_doc.py <spec.txt> <paid_run.txt> <out.docx> <제목> <부제>')
    build(*sys.argv[1:])
