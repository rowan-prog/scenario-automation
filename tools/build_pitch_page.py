#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_pitch_page.py — 피칭 페이지(S급 제작 결정 미팅) 마크다운 정본 → 사내 제출용 docx.

규격 = config/60_pitch_page_standard.md · 절차 = prompts/phase_s_pitch_page.md · skill /pitch-page.

서식(여백·폰트·색상·pt·행간·헬퍼 구조)은 전부 실물 빌더에서 그대로 승계했다:
  projects/17_son_of_the_lamp/_build_pitch_deck_v2.py
  → 「거지 알라딘과 요술램프_피칭페이지_v4.docx」(2026-08-13 미팅 제출본)
새 서식 상수를 발명하지 않는다. 저 파일이 서식 정본이다.

빌드 전에 tools/pitch_page_lint.py 를 호출해 하드 게이트를 건다(--force 로만 우회).
하드 = F1 필드 공란 / F2 회차 수 불일치 / F3 필드 자리 뒤바뀜 / U1 "페이월" 단어 / U2 회차 부제.

사용:
  python tools/build_pitch_page.py <원고.md>
  python tools/build_pitch_page.py <원고.md> --version 2 --confluence
  python tools/build_pitch_page.py <원고.md> --out "C:\\...\\제목_피칭페이지_v2.docx" --force
"""
import argparse

import os
import re
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
# pitch_page_lint 가 import 시점에 sys.stdout 을 UTF-8 로 래핑한다 — 여기서 또 감싸면
# 래퍼 두 겹이 같은 buffer 를 잡아 하나가 닫히면서 I/O 오류가 난다. 저쪽 래핑만 쓴다.
import pitch_page_lint as lint  # noqa: E402

# ── 서식 상수 (_build_pitch_deck_v2.py 승계 — 변경 금지) ────────────────────────
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
ACC = RGBColor(0x8A, 0x5A, 0x10)
LINE = RGBColor(0xC9, 0xB8, 0x92)
FONT = "맑은 고딕"

MARGIN_LR = Cm(2.2)
MARGIN_TB = Cm(1.9)

HARD_GATES = ["F1", "F2", "F3", "U1", "U2"]

# 사내 양식 원문 라벨 (--confluence) — 전각 괄호·전각 공백·병기 언어 원문 그대로. 교정하지 않는다.
CONFLUENCE_LABELS = {
    "title": "타이틀 (Titles, タイトル）",
    "cm": "담당 CM (担当CM名)",
    "basic": "기본정보(Basic Information)",
    "pitch": "해당 IP 피칭 사유 (Core Expectation)",
    "characters": "주요 인물 소개",
    "plot": "줄거리",
    "treatment": "초반 회차 트리트먼트 (1화~Paywall 화)",
}
SECTION_ORDER = ["title", "cm", "basic", "pitch", "characters", "plot", "treatment"]

EP_HEAD_RE = re.compile(r'^\s*(제?\s?\d{1,3}\s*화|EP\.?\s?\d{1,3})\s*$')
CLIFF_RE = re.compile(r'^\s*(?:\*\*)?(클리프|엔딩)(?:\*\*)?\s*[:：]\s*(.+)$')
BOLD_LINE_RE = re.compile(r'^\s*(?:\*\*(.+?)\*\*|###+\s+(.+?))\s*$')
CHAR_HEAD_RE = re.compile(r'^\s*(?:\*\*)?([^\s*].{0,60}?)(?:\*\*)?\s*$')


def strip_md(t):
    """인라인 마크다운 제거 — 본문 산문에 볼드를 만들지 않는다(표준 G14)."""
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)', r'\1', t)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    return t.strip()


class Doc:
    def __init__(self):
        self.d = docx.Document()
        s = self.d.sections[0]
        s.left_margin = s.right_margin = MARGIN_LR
        s.top_margin = s.bottom_margin = MARGIN_TB
        self.n_para = 0
        self.n_bold = 0

    def para(self, space_after=0, space_before=0, keep=False, indent=0, spacing=1.30):
        p = self.d.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(space_before)
        pf.line_spacing = spacing
        pf.keep_with_next = keep
        if indent:
            pf.left_indent = Pt(indent)
        self.n_para += 1
        return p

    def run(self, p, text, *, size=11, color=BLACK, bold=False):
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bold
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        if bold:
            self.n_bold += 1
        return r

    def section(self, title, en=None):
        p = self.para(space_before=18, space_after=7, keep=True)
        self.run(p, title, size=13, bold=True, color=ACC)
        if en:
            self.run(p, "  " + en, size=9.5, color=GREY)

    def body(self, text, space_after=6, indent=0, size=11, bold=False, color=BLACK):
        p = self.para(space_after=space_after, indent=indent)
        self.run(p, text, size=size, bold=bold, color=color)

    def subhead(self, text):
        p = self.para(space_before=10, space_after=3, keep=True)
        self.run(p, text, size=11, bold=True, color=ACC)

    def save(self, path):
        self.d.save(path)


def parse_md(path):
    raw = open(path, encoding="utf-8").read()
    sections = lint.parse_sections(raw)
    return raw, sections


def section_lines(sections, key):
    return [l.rstrip() for l in lint.sec(sections, key).splitlines()]


def korean_title(sections):
    for l in section_lines(sections, "title"):
        m = re.search(r'한국어\s*\(?Korean\)?\s*[:：]\s*(.+)$', strip_md(l))
        if m:
            return m.group(1).strip()
    return None


def english_title(sections):
    for l in section_lines(sections, "title"):
        m = re.search(r'영어\s*\(?English\)?\s*[:：]\s*(.+)$', strip_md(l))
        if m:
            return m.group(1).strip()
    return None


def parse_treatment(sections):
    """[(회차 헤더, 본문 줄 리스트, 클리프 문장 또는 None)] — 원문 순서."""
    eps, cur = [], None
    for raw in section_lines(sections, "treatment"):
        s = strip_md(re.sub(r'^\s*[-*]\s+', '', raw))
        if not s:
            continue
        if EP_HEAD_RE.match(s):
            if cur:
                eps.append(cur)
            cur = [s, [], None]
            continue
        if cur is None:
            continue
        m = CLIFF_RE.match(s)
        if m:
            cur[2] = m.group(2).strip()
        else:
            cur[1].append(s)
    if cur:
        eps.append(cur)
    return eps


def parse_characters(sections):
    """[(헤더, [문단])] — 볼드 줄 또는 `이름 (…) · 역할` 패턴을 헤더로."""
    out, cur = [], None
    for raw in section_lines(sections, "characters"):
        if not raw.strip():
            continue
        line = re.sub(r'^\s*[-*]\s+', '', raw).strip()
        is_head = bool(re.match(r'^\*\*.+\*\*$', line)) or bool(
            re.match(r'^[^\s].{0,50}?\s*[·—-]\s*\S', line) and len(line) <= 70 and '"' not in line
        )
        s = strip_md(line)
        if is_head and s:
            if cur:
                out.append(cur)
            cur = [s, []]
        elif cur is not None:
            cur[1].append(s)
        elif s:
            cur = [s, []]
    if cur:
        out.append(cur)
    return out


def render(sections, confluence, lint_results):
    doc = Doc()
    kr = korean_title(sections) or "(제목 미기입)"
    en = english_title(sections) or ""

    # ── 표지 ──
    t = doc.para(space_after=3)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.run(t, kr, size=20, bold=True)
    if en:
        t2 = doc.para(space_after=4)
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.run(t2, en, size=12, color=ACC)
    t3 = doc.para(space_after=4)
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.run(t3, "─" * 20, size=8, color=LINE)

    def label(key, fallback):
        return CONFLUENCE_LABELS[key] if confluence else fallback

    # ── 담당 CM ──
    cm = [strip_md(l) for l in section_lines(sections, "cm") if l.strip()]
    if cm:
        doc.section(label("cm", "담당 CM"))
        doc.body(" ".join(cm))

    # ── 기본정보 (로그라인 본문만 볼드 — 표준 G14) ──
    basic = [l for l in section_lines(sections, "basic") if l.strip()]
    if basic:
        doc.section(label("basic", "기본정보"), "Basic Information" if not confluence else None)
        in_logline = False
        for raw in basic:
            s = strip_md(re.sub(r'^\s*[-*]\s+', '', raw))
            if not s:
                continue
            if re.match(r'^(로그라인|Logline)', s, re.I):
                in_logline = True
                head, _, rest = re.split(r'([:：])', s, maxsplit=1)[0], None, ""
                m = re.match(r'^(로그라인[^:：]*|Logline[^:：]*)[:：]\s*(.*)$', s, re.I)
                if m:
                    p = doc.para(space_after=3)
                    doc.run(p, m.group(1).strip() + "  ", size=10.5, bold=True, color=GREY)
                    if m.group(2).strip():
                        doc.run(p, m.group(2).strip(), size=11, bold=True)
                else:
                    doc.body(s, space_after=3, size=10.5, bold=True, color=GREY)
                continue
            if lint.LOGLINE_LABEL_RE.match(s) is None and any(
                s.startswith(lbl) for lbl in lint.FIELD_LABELS
            ):
                in_logline = False
            doc.body(s, space_after=3, bold=in_logline and not any(
                s.startswith(lbl) for lbl in lint.FIELD_LABELS))

    # ── 피칭 사유 (소제목만 볼드) ──
    pitch = [l for l in section_lines(sections, "pitch") if l.strip()]
    if pitch:
        doc.section(label("pitch", "해당 IP 피칭 사유"),
                    "Core Expectation" if not confluence else None)
        for raw in pitch:
            line = re.sub(r'^\s*[-*]\s+', '', raw).strip()
            s = strip_md(line)
            if not s:
                continue
            is_sub = bool(BOLD_LINE_RE.match(line)) and not s.endswith(".")
            if is_sub:
                doc.subhead(s)
            else:
                doc.body(s)

    # ── 인물 ──
    chars = parse_characters(sections)
    if chars:
        doc.section(label("characters", "주요 인물 소개"))
        for head, paras in chars:
            doc.subhead(head)
            for i, t_ in enumerate(paras):
                doc.body(t_, space_after=0 if i == len(paras) - 1 else 4)

    # ── 줄거리 ──
    plot = [strip_md(l) for l in section_lines(sections, "plot") if l.strip()]
    if plot:
        doc.section(label("plot", "줄거리"))
        for t_ in plot:
            doc.body(t_)

    # ── 트리트먼트 (마지막 무료 화 클리프만 볼드 — 표준 G14) ──
    eps = parse_treatment(sections)
    if eps:
        doc.section(label("treatment", "초반 회차 트리트먼트"),
                    "1화~무료 마지막 화" if not confluence else None)
        last = len(eps) - 1
        for i, (head, lines, cliff) in enumerate(eps):
            h = doc.para(space_before=10, space_after=3, keep=True)
            doc.run(h, head, size=11, bold=True, color=ACC)
            for t_ in lines:
                doc.body(t_, space_after=3)
            if cliff:
                p = doc.para(space_after=0)
                doc.run(p, "클리프  ", size=10.5, bold=True, color=GREY)
                doc.run(p, cliff, size=10.5, bold=(i == last))
    return doc, kr, len(eps), len(chars)


def main():
    ap = argparse.ArgumentParser(description="피칭 페이지 md → docx 빌더")
    ap.add_argument("src", help="피칭 페이지 마크다운 정본 (_02_pitch_page.md)")
    ap.add_argument("--version", default="1", help="산출 파일 버전 번호 (기본 1)")
    ap.add_argument("--confluence", action="store_true",
                    help="섹션 제목을 사내 양식 원문 라벨로 렌더")
    ap.add_argument("--force", action="store_true", help="하드 게이트 FAIL이어도 빌드")
    ap.add_argument("--out", default=None, help="산출 docx 경로 지정")
    a = ap.parse_args()

    src = os.path.abspath(a.src)
    if not os.path.exists(src):
        print(f"입력 파일 없음: {src}")
        return 2

    raw, sections = parse_md(src)

    # ── 게이트 ──
    # run_all_gates → [(게이트 id, 게이트명, 판정, 상세), ...]
    results = lint.run_all_gates(sections, raw)
    verdicts = {gid: verdict for gid, _name, verdict, _detail in results}
    hard_fail = [g for g in HARD_GATES if verdicts.get(g) == "FAIL"]
    n_fail = sum(1 for v in verdicts.values() if v == "FAIL")
    n_warn = sum(1 for v in verdicts.values() if v == "WARN")

    print(f"게이트: FAIL {n_fail} · WARN {n_warn}  (전체 판정 = python tools/pitch_page_lint.py \"{src}\")")
    for gid, name, verdict, detail in results:
        if verdict in ("FAIL", "WARN"):
            mark = "🚨" if gid in HARD_GATES and verdict == "FAIL" else "  "
            print(f"  {mark}[{verdict}] {gid} {name} — {detail}")

    if hard_fail and not a.force:
        print(f"\n빌드 중단 — 하드 게이트 FAIL: {', '.join(hard_fail)}")
        print("원고를 고치고 다시 실행하라. 그대로 빌드하려면 --force.")
        return 1
    if hard_fail:
        print(f"\n⚠️ --force — 하드 게이트 FAIL을 무시하고 빌드: {', '.join(hard_fail)}")

    doc, kr, n_ep, n_char = render(sections, a.confluence, results)

    out = a.out or os.path.join(os.path.dirname(src), f"{kr}_피칭페이지_v{a.version}.docx")
    doc.save(out)
    print(f"\nOK docx  {out}")
    print(f"  섹션 {sum(1 for k in SECTION_ORDER if lint.sec(sections, k).strip())}종 · "
          f"인물 {n_char}인 · 회차 {n_ep}화 · 문단 {doc.n_para} · 볼드 런 {doc.n_bold}")
    print(f"  라벨 = {'사내 양식 원문(--confluence)' if a.confluence else 'md 원문'} · "
          f"마지막 무료 화 클리프만 볼드(G14)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
