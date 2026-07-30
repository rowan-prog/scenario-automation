# -*- coding: utf-8 -*-
"""Vigloo(플랫폼) 기획안 docx 빌더 — 템플릿 원본을 복사해 셀을 직접 채운다 (phase_p · 2026-07-28 전면 개정)

구버전(자체 섹션 레이아웃 생성)은 템플릿 불일치 판정으로 폐기.
정본 템플릿 = config/vigloo_template/Vigloo_Proposal_Template_원본.docx (단일 표 51행×2열).
이 빌더는 템플릿 docx를 열어 지정 행의 우측 셀만 채우고, 안 쓰는 행을 삭제해 저장한다
— 라벨·서식·표 구조는 템플릿 원본이 그대로 보존된다.

사용: python tools/build_vigloo_proposal.py <spec.txt> <out.docx>

spec 형식 (UTF-8 · 행 번호 = 원본 템플릿 표의 0-기준 행 번호):
  #TEMPLATE <경로>          (선택 · 기본 = config/vigloo_template/Vigloo_Proposal_Template_원본.docx)
  #DOCTITLE <문서 최상단 타이틀>
  #FILL <행번호>
  <우측 셀 내용 줄들... 다음 지시어 전까지. 빈 줄 = 빈 문단으로 보존>
  #FILLLABEL <행번호>
  <좌측 라벨 셀을 교체할 때만 (기본 = 원문 유지 — 임의 라벨 신설·삭제 금지 원칙)>
  #DELETE 16,17,18,20      (미사용 행 삭제 · 쉼표 구분 · 원본 행 번호 기준)

참고(원본 템플릿 행 지도): 0 기본정보헤더 / 1 타이틀 / 2 담당자 / 3 제작형태 / 4 장르
5 키워드 / 6 로그라인 / 7 총회차 / 8 타깃 / 9 AI Type / 10 레퍼런스 / 11 시놉헤더 / 12 줄거리
13 캐릭터헤더 / 14 주요인물헤더 / 15~18 남주 / 19~22 여주 / 23 서브헤더 / 24~27 서브여주
28~31 서브남주 / 32 (빈 행) / 33 핵심아이디어헤더 / 34 사이다훅 / 35 세계관 / 36 내용요건헤더
37 필수 / 38 금기 / 39 트리트먼트헤더 / 40 회별클리프 / 41 제작적절성헤더 / 42 제작비
43 납품기일 / 44 릴리즈 / 45 샘플헤더 / 46 샘플회차 / 47 샘플납품일 / 48 샘플비용
49 검토자료헤더 / 50 각색대본
"""
import sys, os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DEFAULT_TEMPLATE = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                'config', 'vigloo_template', 'Vigloo_Proposal_Template_원본.docx')


def parse_spec(path):
    template = DEFAULT_TEMPLATE
    doctitle = None
    fills = {}       # row_idx -> [lines]
    fill_labels = {} # row_idx -> [lines]
    deletes = []
    cur = None
    for ln in open(path, encoding='utf-8').read().split('\n'):
        if ln.startswith('#TEMPLATE '):
            template = ln[10:].strip(); cur = None
        elif ln.startswith('#DOCTITLE '):
            doctitle = ln[10:].strip(); cur = None
        elif ln.startswith('#FILLLABEL '):
            idx = int(ln[11:].strip()); cur = fill_labels.setdefault(idx, [])
        elif ln.startswith('#FILL '):
            idx = int(ln[6:].strip()); cur = fills.setdefault(idx, [])
        elif ln.startswith('#DELETE '):
            deletes += [int(x) for x in ln[8:].replace(' ', '').split(',') if x]; cur = None
        elif cur is not None:
            cur.append(ln)
    strip_tail = lambda lines: [l for i, l in enumerate(lines)
                                if any(x.strip() for x in lines[i:])]  # 꼬리 빈 줄 제거
    return template, doctitle, {k: strip_tail(v) for k, v in fills.items()}, \
        {k: strip_tail(v) for k, v in fill_labels.items()}, sorted(set(deletes), reverse=True)


def own_cell(row, col_idx):
    # 세로 병합(vMerge) 해제 후 그 행 소유의 독립 셀 반환 — 템플릿 캐릭터 4행이
    # 한 칸으로 병합돼 있어, 병합을 안 풀면 뒤에 채운 인물이 앞 인물을 덮어쓴다.
    from docx.table import _Cell
    tc = row._tr.tc_lst[col_idx]
    tcPr = tc.tcPr
    if tcPr is not None:
        for v in tcPr.findall(qn('w:vMerge')):
            tcPr.remove(v)
    return _Cell(tc, row)


def set_cell_text(cell, lines):
    # 첫 문단만 남기고 비운 뒤 줄당 문단 1개. 폰트 = 맑은 고딕 10pt(한중 공통) 명시.
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    paras = [first]
    for _ in range(len(lines) - 1):
        paras.append(cell.add_paragraph())
    for p, text in zip(paras, lines):
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')


def build(spec_path, out_path):
    template, doctitle, fills, fill_labels, deletes = parse_spec(spec_path)
    doc = Document(template)
    if doctitle is not None:
        p = doc.paragraphs[0]
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        run = p.add_run(doctitle)
        run.bold = True
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    table = doc.tables[0]
    rows = list(table.rows)  # 원본 행 번호 기준 스냅샷
    for idx, lines in fill_labels.items():
        set_cell_text(own_cell(rows[idx], 0), lines)
    for idx, lines in fills.items():
        own_cell(rows[idx], 0)  # 라벨 열 병합도 해제 (내용은 원문 유지)
        set_cell_text(own_cell(rows[idx], 1), lines)
    for idx in deletes:
        if idx in fills or idx in fill_labels:
            raise SystemExit(f'행 {idx}: FILL과 DELETE 동시 지정 — spec 오류')
        table._tbl.remove(rows[idx]._tr)
    doc.save(out_path)
    kept = len(list(Document(out_path).tables[0].rows))
    print(f'OK: {out_path} (rows {len(rows)} -> {kept}, filled {len(fills)})')


if __name__ == '__main__':
    if len(sys.argv) != 3:
        raise SystemExit('사용: python tools/build_vigloo_proposal.py <spec.txt> <out.docx>')
    build(sys.argv[1], sys.argv[2])
