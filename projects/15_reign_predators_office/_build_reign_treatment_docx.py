# -*- coding: utf-8 -*-
"""군림 회차 트리트먼트 md → docx (작가 전달용)"""
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT = 'Malgun Gothic'


def build(md_path, out_path):
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)

    def para(space_after=6, space_before=0, indent=0.0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    def run(p, text, bold=False, size=10.5):
        r = p.add_run(text)
        r.bold = bold
        r.font.name = FONT
        r.font.size = Pt(size)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        return r

    lines = open(md_path, encoding='utf-8').read().split('\n')
    ep_count = 0

    for line in lines:
        t = line.strip()
        if not t or t == '---':
            continue

        if t.startswith('# '):
            run(para(space_after=18), t[2:], bold=True, size=16)
            continue

        if t.startswith('### '):
            ep_count += 1
            run(para(space_after=6, space_before=14), t[4:], bold=True, size=12)
            continue

        if t.startswith('## '):
            ep_count += 1
            if ep_count > 1:
                run(para(space_after=0), '').add_break(WD_BREAK.PAGE)
            run(para(space_after=10, space_before=0), t[3:], bold=True, size=13)
            continue

        if t.startswith('> '):
            body = t[2:].strip()
            m = re.match(r'\*\*(.+?)\*\*\s*(.*)', body)
            p = para(space_after=4, indent=0.35)
            if m:
                run(p, m.group(1), bold=True)
                run(p, '\t' + m.group(2).strip())
            else:
                run(p, body.replace('**', ''))
            continue

        run(para(space_after=8), t.replace('**', ''))

    doc.save(out_path)
    print(f'OK: {out_path} / 회차 블록 {ep_count}')


if __name__ == '__main__':
    build(sys.argv[1], sys.argv[2])
