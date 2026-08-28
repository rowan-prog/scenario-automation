# -*- coding: utf-8 -*-
"""군림 회차 트리트먼트 docx v3 — free_run(N화 + 한 줄=한 비트) + paid_run([K…] 라벨 + NN화|본문)"""
import re, sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def build(free_path, paid_path, out_path, title, subtitle):
    doc = Document()
    def para(text, bold=False, size=10, space=6, italic=False):
        p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.name = 'Malgun Gothic'; r.font.size = Pt(size)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        p.paragraph_format.space_after = Pt(space); return p
    para(title, bold=True, size=15, space=2); para(subtitle, size=9, space=14)
    para('무료회차 (1~8화)', bold=True, size=12, space=8)
    n_free = 0
    for line in open(free_path, encoding='utf-8').read().split('\n'):
        t = line.strip()
        if not t: continue
        if re.fullmatch(r'\d+화', t): para(t, bold=True, size=11, space=6); n_free += 1
        else: para(t, size=10, space=2)
    para('', space=10)
    para('유료회차 (9~50화)', bold=True, size=12, space=8)
    n_paid = 0
    for row in open(paid_path, encoding='utf-8').read().split('\n'):
        t = row.strip()
        if not t: continue
        if t.startswith('['): para(t, bold=True, italic=True, size=10, space=6); continue
        ep, body = t.split('|', 1); para(ep, bold=True, size=11, space=2); para(body, size=10, space=10); n_paid += 1
    doc.save(out_path); print(f'OK: {out_path} / 무료 {n_free}화 · 유료 {n_paid}화')

if __name__ == '__main__': build(*sys.argv[1:6])
