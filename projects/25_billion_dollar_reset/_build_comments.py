# -*- coding: utf-8 -*-
"""작가 발송용 검수 코멘트 docx 빌더.
마크업:  > 원문(中)   >~ 원문 한국어 대역   >> 교체안(中)   >>~ 교체안 한국어 대역
         → 판정/방향   ## 대제목   ### 앵커   --- 구분선
"""
import io, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_LINE_SPACING

OUT_DIR  = r'C:\Users\Rowan\scenario-automation\projects\25_billion_dollar_reset'
SRC_MD   = os.path.join(OUT_DIR, 'MY BILLION DOLLAR RESET_1-50화_검수코멘트_v1.md')
OUT_DOCX = os.path.join(OUT_DIR, 'MY BILLION DOLLAR RESET_1-50화_검수코멘트_v1.docx')

GREY_CN = (0x44, 0x44, 0x44)   # 원문 중국어
GREY_KO = (0x8A, 0x8A, 0x8A)   # 원문 대역
RED_CN  = (0xA8, 0x1F, 0x1F)   # 교체안 중국어
RED_KO  = (0xC2, 0x6A, 0x6A)   # 교체안 대역
BLUE    = (0x14, 0x4A, 0x8A)
RULE    = (0xC8, 0xC8, 0xC8)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(2.0)

n = doc.styles['Normal']
n.font.name = '맑은 고딕'
n.font.size = Pt(10)
pf = n.paragraph_format
pf.space_after = Pt(1)
pf.space_before = Pt(0)
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.05


def para(text, *, size=10, color=None, bold=False, indent=0.0, before=0, after=1,
         spacing=1.05, mono=False):
    p = doc.add_paragraph()
    f = p.paragraph_format
    f.left_indent = Cm(indent)
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = spacing
    text = text.replace('`', '')
    for i, chunk in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.bold = bold or (i % 2 == 1)
        r.font.size = Pt(size)
        if mono:
            r.font.name = 'Malgun Gothic'
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


for raw in io.open(SRC_MD, encoding='utf-8').read().split('\n'):
    s = raw.rstrip()
    t = s.strip()

    if not t:
        para('', size=4, after=0)
        continue

    if t.startswith('# '):
        para(t[2:], size=17, bold=True, color=BLUE, after=3)
        continue

    if t.startswith('## '):
        para(t[3:], size=13.5, bold=True, color=BLUE, before=13, after=4)
        continue

    if t.startswith('### '):
        para(t[4:], size=11, bold=True, before=8, after=3)
        continue

    if t.startswith('>>~'):
        para(t[3:].strip(), size=9, color=RED_KO, indent=1.5, after=3)
        continue
    if t.startswith('>>'):
        para(t[2:].strip(), size=10, color=RED_CN, indent=1.0, after=0)
        continue
    if t.startswith('>~'):
        para(t[2:].strip(), size=9, color=GREY_KO, indent=1.5, after=3)
        continue
    if t.startswith('>'):
        para(t[1:].strip(), size=10, color=GREY_CN, indent=1.0, after=0)
        continue

    if t.startswith('→'):
        para(t, indent=0.25, before=3, after=3)
        continue

    if t.startswith('- '):
        para('·  ' + t[2:], indent=0.4, after=2)
        continue

    if t == '---':
        para('─' * 62, size=8, color=RULE, before=8, after=8)
        continue

    para(t, spacing=1.25, after=3)

doc.save(OUT_DOCX)
print('saved:', OUT_DOCX)
