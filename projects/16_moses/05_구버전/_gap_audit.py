# -*- coding: utf-8 -*-
import os as _os; _os.chdir(r'C:\Users\Rowan\scenario-automation\projects\16_moses')
"""작가 초고(_draft_paras.txt) 전수 스캔 → 수정 지시서가 실제로 덮었는지 대조.
   지시서가 커버 못 한 위반 단락 = 구멍."""
import io, re, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

BASE = r"C:\Users\Rowan\scenario-automation\projects\16_moses" + "\\"
FB = BASE + '내 남편은 거지 모세_1-21화 수정 지시서_v1.docx'

# ── 초고 로드 ──────────────────────────────────────────
draft = {}
for ln in open(BASE + '03_작업파일/_draft_paras.txt', encoding='utf-8'):
    if '\t' not in ln:
        continue
    n, t = ln.rstrip('\n').split('\t', 1)
    draft[int(n)] = t

E = {1: 0, 2: 84, 3: 138, 4: 184, 5: 232, 6: 285, 7: 328, 8: 378, 9: 442,
     10: 512, 11: 568, 12: 632, 13: 691, 14: 729, 15: 788, 16: 849,
     17: 893, 18: 958, 19: 1011, 20: 1054, 21: 1098}


def ep_of(n):
    cur = 1
    for k in sorted(E):
        if n >= E[k]:
            cur = k
    return cur


# 재작성 구간 = 지시서가 실물 대본으로 통째 대체
REWRITE = list(range(512, 568)) + list(range(893, 1137))

# ── 지시서가 인용한 단락 번호 수집 ────────────────────
d = Document(FB)
cited = set()
cite_re = re.compile(r'\d{1,4}')
for t in d.tables:
    for r in t.rows:
        head = r.cells[0].text
        for m in cite_re.findall(head):
            cited.add(int(m))
        # 일괄 치환 블록은 '찾을 말' 칸에 번호가 들어있음
        if len(r.cells) >= 2:
            for m in cite_re.findall(r.cells[1].text):
                cited.add(int(m))
for p in d.paragraphs:
    if '삭제' in p.text or '전량' in p.text or '단락' in p.text:
        for m in cite_re.findall(p.text):
            cited.add(int(m))
# 3화 145~151 일괄 삭제
cited |= set(range(145, 152))

# ── 위반 패턴 ────────────────────────────────────────
RULES = [
    ('G27 제국', r'제국'),
    ('G27 집정관', r'집정관'),
    ('G27 평의회', r'평의회'),
    ('G27 원로', r'원로'),
    ('G27 정표', r'정표'),
    ('G27 성물', r'성물'),
    ('G27 의원', r'의원'),
    ('G27 지옥', r'지옥'),
    ('G27 법정', r'법정'),
    ('G27 왕비 서임', r'서임'),
    ('G28 연꽃/화관', r'연꽃|화관|수문'),
    ('G23 황금갑옷/수호대', r'황금 갑옷|황금빛 갑옷|수호대|정예'),
    ('G18 복수형 신들', r'신들'),
    ('G10/G15 몸에서 빛', r'몸에서 (강렬한 )?황금|신력이 일렁|황금빛 신력|두 눈이 섬뜩한 황금'),
    ('G11 사형선고/강림', r'죽음을 피하지 못할|강림|옥좌에 앉아 있는 모세'),
    ('G8 비행/모래폭풍 이동', r'모래폭풍을 타고|하늘로 비상'),
    ('G7 피가 돌을 태움', r'피가 닿은 돌바닥이'),
    ('G14 펜던트=배터리', r'신력을 지켜주'),
    ('G13 지팡이=통로', r'지팡이를 쥐고 바닥을 쾅|지팡이에서 신력'),
    ('G24 성소=보물고', r'성소에서 온 보물|신성한 봉헌물'),
]

print('=' * 72)
print('작가 초고 위반 스캔 → 지시서 커버 여부')
print('=' * 72)
gaps = []
for name, pat in RULES:
    rx = re.compile(pat)
    hits = [n for n, t in draft.items() if rx.search(t)]
    if not hits:
        continue
    uncov = [n for n in hits if n not in cited and n not in REWRITE]
    print(f'\n[{name}] 초고 {len(hits)}건 / 미커버 {len(uncov)}건')
    for n in uncov:
        gaps.append((name, n))
        print(f'   ✗ {n}(EP{ep_of(n)}) {draft[n][:95]}')

print()
print('=' * 72)
print(f'미커버 총계: {len(gaps)}건')
