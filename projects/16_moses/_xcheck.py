# -*- coding: utf-8 -*-
"""가이드 v10 / 수정 지시서 v1 / 트리트먼트 v2 3자 금칙어 교차 점검."""
import io, sys, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

BASE = r"C:\Users\Rowan\scenario-automation\projects\16_moses" + "\\"
DOCS = {
    'GUIDE': BASE + '내 남편은 거지 모세_각색 가이드_v10.docx',
    'FEEDBACK': BASE + '내 남편은 거지 모세_1-21화 수정 지시서_v1.docx',
    'TREATMENT': BASE + '내 남편은 거지 모세_회차 트리트먼트_1-50화_v2.docx',
}

BAN = ['제국', '집정관', '평의회', '왕비 서임', '유일신', '진정한 신',
       '정표', '성물', '피시아', '노리라', '제시타', '연꽃', '화관',
       '황금 갑옷', '수호대', '법정', '의원', '지옥', '아레오파고스',
       '올림포스', '아폴론', '레토', '칼리스토', '도리에우스', '델포이']


def dump(path):
    d = Document(path)
    out = []
    for p in d.paragraphs:
        out.append(p.text)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                out.append(c.text)
    return out


for name, path in DOCS.items():
    paras = dump(path)
    print('=' * 60)
    print(name, '— 단락', len(paras))
    for w in BAN:
        hits = [(i, t) for i, t in enumerate(paras) if w in t]
        if hits:
            print(f'  [{w}] {len(hits)}건')
            for i, t in hits[:6]:
                print('     ', t.replace('\n', ' / ')[:150])
