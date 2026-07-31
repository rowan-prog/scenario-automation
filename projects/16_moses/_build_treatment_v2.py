# -*- coding: utf-8 -*-
"""16_moses_p1_treatment_full_v2.md → 회차 트리트먼트 docx (v2)"""
import io, re, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\Rowan\scenario-automation\projects\16_moses\16_moses_p1_treatment_full_v2.md"
OUT = r"C:\Users\Rowan\scenario-automation\projects\16_moses\내 남편은 거지 모세_회차 트리트먼트_1-50화_v2.docx"

doc = Document()
st = doc.styles['Normal']
st.font.name = '맑은 고딕'
st.font.size = Pt(10)
st.element.rPr.rFonts.set(
    __import__('docx').oxml.ns.qn('w:eastAsia'), '맑은 고딕')
st.paragraph_format.space_after = Pt(4)

lines = open(SRC, encoding='utf-8').read().split('\n')

i = 0
tbl_buf = []


def flush_table():
    global tbl_buf
    if not tbl_buf:
        return
    rows = [[c.strip() for c in r.strip().strip('|').split('|')] for r in tbl_buf
            if not re.match(r'^\|[\s:\-|]+\|$', r.strip())]
    tbl_buf = []
    if not rows:
        return
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            v = v.replace('**', '')
            c.text = v
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(8.5)
                    if ri == 0:
                        rr.bold = True


def emit_rich(p, text):
    """**bold** 마크업 처리"""
    for j, seg in enumerate(text.split('**')):
        if not seg:
            continue
        r = p.add_run(seg)
        r.bold = (j % 2 == 1)


while i < len(lines):
    ln = lines[i]
    s = ln.strip()

    if s.startswith('|'):
        tbl_buf.append(s)
        i += 1
        continue
    flush_table()

    if not s:
        i += 1
        continue

    if s.startswith('# '):
        p = doc.add_paragraph()
        r = p.add_run(s[2:])
        r.bold = True
        r.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(14)
    elif s.startswith('## '):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(s[3:])
        r.bold = True
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    elif s == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    elif s.startswith('**클리프행어**') or s.startswith('**엔딩**'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        emit_rich(p, s)
        for r in p.runs:
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(10)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        emit_rich(p, s)
    i += 1

flush_table()
doc.save(OUT)
print('WROTE', OUT)
