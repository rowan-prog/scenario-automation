# -*- coding: utf-8 -*-
import os as _os; _os.chdir(r'C:\Users\Rowan\scenario-automation\projects\16_moses')
"""검토 에이전트용 텍스트 덤프."""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

B = r"C:\Users\Rowan\scenario-automation\projects\16_moses" + "\\"
JOBS = [
    ('내 남편은 거지 모세_각색 가이드_v15.docx', '03_작업파일/_export_guide_v15.txt'),
    ('내 남편은 거지 모세_1-21화 수정 지시서_v3.docx', '03_작업파일/_export_feedback_v1.txt'),
]

for src, dst in JOBS:
    d = Document(B + src)
    out = []
    body = d.element.body
    ti, pi = 0, 0
    tables, paras = d.tables, d.paragraphs
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p' and pi < len(paras):
            t = paras[pi].text
            pi += 1
            if t.strip():
                out.append(t)
        elif tag == 'tbl' and ti < len(tables):
            tb = tables[ti]
            ti += 1
            for r in tb.rows:
                out.append(' | '.join(c.text.replace('\n', ' ') for c in r.cells))
            out.append('')
    open(B + dst, 'w', encoding='utf-8').write('\n'.join(out))
    print('WROTE', B + dst, len('\n'.join(out)), 'chars')
