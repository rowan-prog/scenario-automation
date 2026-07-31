# -*- coding: utf-8 -*-
"""수정 지시서: '이렇게'(교체안) 열 + 실물 대본 블록에만 금칙어가 남았는지 점검."""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

BASE = r"C:\Users\Rowan\scenario-automation\projects\16_moses" + "\\"
FB = BASE + '내 남편은 거지 모세_1-21화 수정 지시서_v1.docx'

BAN = ['제국', '집정관', '평의회', '유일신', '진정한 신', '정표', '성물',
       '피시아', '노리라', '제시타', '연꽃', '화관', '황금 갑옷', '수호대',
       '법정', '의원', '지옥', '올림포스', '아폴론', '레토', '모래폭풍',
       '신전의 정예', '신들이시여', '통째로 집어삼킨다', '기름 먹인']

d = Document(FB)

print('### 교체안(3열 표의 마지막 칸)에 남은 금칙어')
n = 0
for t in d.tables:
    for r in t.rows:
        cells = [c.text for c in r.cells]
        if len(cells) < 2:
            continue
        tgt = cells[-1]
        for w in BAN:
            if w in tgt:
                n += 1
                print(f'  [{w}] {cells[0][:14]} | {tgt[:120]}')
print('  → 없음' if n == 0 else f'  → {n}건')

print()
print('### 실물 대본 블록(△·화자 줄)에 남은 금칙어')
m = 0
for p in d.paragraphs:
    tx = p.text
    if not (tx.startswith('△') or tx.startswith('제1') or tx.startswith('#')
            or (':' in tx and tx.split(':')[0].strip() in
                ('모세', '다말', '델릴라', '파라오', '라반', '네페라', '아론',
                 '재상', '근위대장', '장로1', '마술사1', '마술사2', '간수',
                 '군중1', '모세VO', '다말VO'))):
        continue
    for w in BAN:
        if w in tx:
            m += 1
            print(f'  [{w}] {tx[:130]}')
print('  → 없음' if m == 0 else f'  → {m}건')
