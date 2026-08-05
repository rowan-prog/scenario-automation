# -*- coding: utf-8 -*-
import re, docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

CN_RAW = open('17_son_of_the_lamp_source_CN.md', encoding='utf-8').read().split('\n')
KR_RAW = open('17_son_of_the_lamp_source_KR.md', encoding='utf-8').read().split('\n')

cn = [l.strip() for l in CN_RAW if l.strip()]
kr = [l.strip() for l in KR_RAW if l.strip() and not l.strip().startswith('[^') and l.strip() != '---']

# --- fix original line-wrap artifact: merge continuation line into previous dialogue
for i, l in enumerate(cn):
    if l == '您一定要给我出这口恶气！':
        cn[i-1] = cn[i-1] + ' ' + l
        cn.pop(i); break

ci = cn.index('第1集'); ki = kr.index('# 제1화')
cn_front, cn_body = cn[:ci], cn[ci:]
kr_front, kr_body = kr[:ki], kr[ki:]

# ---------- front matter: pair by markers ----------
MARK = [('世界观设定：','## 세계관 설정'), ('核心人物小传：','## 핵심 인물 소개'),
        ('阿拉丁(Aladdin)','**알라딘(Aladdin)**'), ('索拉娅(Soraya)','**소라야(Soraya) — 여주인공**'),
        ('吉尼(Genie)','**지니(Genie) — 최고의 남자 조연/요술램프의 정령**'),
        ('马利克(Malik)','**말리크(Malik) — 메인 빌런**'), ('卡西姆(Qasim)','**카심(Qasim) — 귀족 공자/연적**'),
        ('苏丹(Sultan)','**술탄(Sultan) — 권위 캐릭터**'), ('集纲：','## 집강 (회차 개요)')]
for n in range(1, 8):
    MARK.append(('第%d集'%n, '**제%d집**'%n))

def split_blocks(lines, keys):
    idx = []
    for k in keys:
        found = [i for i, l in enumerate(lines) if l.startswith(k)]
        idx.append(found[0] if found else None)
    blocks = []
    for j, s in enumerate(idx):
        if s is None: blocks.append([]); continue
        nxt = next((idx[m] for m in range(j+1, len(idx)) if idx[m] is not None), len(lines))
        blocks.append(lines[s:nxt])
    return blocks

cb = split_blocks(cn_front, [m[0] for m in MARK])
kb = split_blocks(kr_front, [m[1] for m in MARK])

# ---------- body: line-by-line ----------
assert len(cn_body) == len(kr_body), (len(cn_body), len(kr_body))

# ---------- doc ----------
d = docx.Document()
sec = d.sections[0]
sec.left_margin = sec.right_margin = Cm(2.0); sec.top_margin = sec.bottom_margin = Cm(1.8)
GREY = RGBColor(0x7A, 0x7A, 0x7A); BLACK = RGBColor(0x1A, 0x1A, 0x1A)
IDX = RGBColor(0xB0, 0xB0, 0xB0); ACC = RGBColor(0x8A, 0x5A, 0x10)

def para(space_after=0, space_before=0, keep=False, indent=0):
    p = d.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    pf.line_spacing = 1.18; pf.keep_with_next = keep
    if indent: pf.left_indent = Pt(indent)
    return p

def run(p, text, *, cnfont=False, size=11, color=BLACK, bold=False):
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold
    name = '微软雅黑' if cnfont else '맑은 고딕'
    r.font.name = name; r._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    return r

MD = re.compile(r'(\*\*[^*]+\*\*)')
def kr_line(p, text, size=11, color=BLACK):
    text = re.sub(r'\[\^\d\]', '', text)
    for seg in MD.split(text):
        if not seg: continue
        run(p, seg[2:-2] if seg.startswith('**') else seg, size=size, color=color, bold=seg.startswith('**'))

def pair(k_text, c_text, n=None):
    p = para(space_after=1, keep=True)
    if n is not None:
        run(p, '%02d  ' % n, size=8, color=IDX)
    kr_line(p, k_text)
    p2 = para(space_after=9, indent=(20 if n is not None else 0))
    run(p2, re.sub(r'^\*\*|\*\*', '', c_text), cnfont=True, size=9.5, color=GREY)

# ----- title -----
t = para(space_after=2); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t, '거지 알라딘과 요술램프  ·  무료회차 1~7화', size=17, bold=True)
t2 = para(space_after=4); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t2, '《乞丐阿拉丁与神灯》 第1~7集', cnfont=True, size=12, color=GREY)
t3 = para(space_after=20); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t3, '한국어 / 中文 원문 대조 합본  —  윗줄 = 한국어(읽기용), 아랫줄 = 중국어 원문(인용용)', size=9, color=GREY)

# ----- front matter -----
for (ck, kk), cblk, kblk in zip(MARK, cb, kb):
    if not cblk and not kblk: continue
    head_kr = kblk[0] if kblk else ''
    if head_kr.startswith('## '):
        h = para(space_before=10, space_after=6, keep=True)
        run(h, head_kr[3:], size=13, bold=True, color=ACC)
        run(h, '   ' + cblk[0], cnfont=True, size=10, color=GREY)
        cblk, kblk = cblk[1:], kblk[1:]
    elif head_kr:
        h = para(space_before=10, space_after=4, keep=True)
        run(h, re.sub(r'\*\*', '', head_kr), size=11.5, bold=True)
        run(h, '   ' + cblk[0], cnfont=True, size=9.5, color=GREY)
        cblk, kblk = cblk[1:], kblk[1:]
    for l in kblk:
        p = para(space_after=1); kr_line(p, l)
    for l in cblk:
        p = para(space_after=8); run(p, l, cnfont=True, size=9.5, color=GREY)

# ----- body -----
i = 0; scene_n = 0
while i < len(kr_body):
    k, c = kr_body[i], cn_body[i]
    if k.startswith('# '):                      # episode
        d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = para(space_after=10, keep=True)
        run(h, k[2:], size=16, bold=True)
        run(h, '   ' + c, cnfont=True, size=11, color=GREY)
        scene_n = 0
    elif k.startswith('### '):                  # scene
        h = para(space_before=12, space_after=2, keep=True)
        run(h, k[4:], size=12.5, bold=True, color=ACC)
        h2 = para(space_after=6, keep=True)
        run(h2, c, cnfont=True, size=10, color=GREY)
        scene_n = 0
    else:
        scene_n += 1
        pair(k, c, scene_n)
    i += 1

# ----- 역주 -----
d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
h = para(space_after=8); run(h, '역주 (원문 오류)', size=13, bold=True, color=ACC)
NOTES = [
 '집강 2집 「阿拉丁有任何多余的动作」 — 부정어(没) 누락. 문맥상 “군더더기 동작 하나 없이”.',
 '집강 7집 「不够他安抚住卡西姆后」 — 오타, 문맥상 「不过」.',
 '001-2 자막 「苏曼」 — 인물명은 「苏丹」(술탄).',
 '004-1 사건 장소는 大殿인데 지시는 「封锁寝宫」.',
 '007-2 「但还是选择卡西姆。」 — 서술어 누락(安抚 등). 번역은 문맥으로 보충.',
 '003-2 「解决了哪些渣滓」 — 「那些」의 오타.',
 '001-2 「卡尔马的王座」 — 국명은 「卡马尔」. 「索拉雅」/「索拉娅」 혼용(001-2·002-2).',
]
for n, t_ in enumerate(NOTES, 1):
    p = para(space_after=4); run(p, '%d. ' % n, size=10, color=GREY); kr_line(p, t_, size=10)

out = '거지 알라딘과 요술램프_무료회차 1-7화_한중 대조본.docx'
d.save(out)
print('OK', out, '| body pairs', len(kr_body), '| front blocks', len(MARK))
