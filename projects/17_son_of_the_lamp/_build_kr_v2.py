# -*- coding: utf-8 -*-
"""작가 수정고(2026-08-05) 한국어 번역본 + 한중 대조본 docx 빌더."""
import re, docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

cn = [l.strip() for l in open('17_son_of_the_lamp_source_CN_v2.md', encoding='utf-8') if l.strip()]
kr = [l.strip() for l in open('17_son_of_the_lamp_source_KR_v2.md', encoding='utf-8')
      if l.strip() and not l.strip().startswith('[^') and l.strip() != '---']

# 원문 줄바꿈 아티팩트: 이어지는 대사 한 줄을 앞줄에 병합
for i, l in enumerate(cn):
    if l == '您一定要给我出这口恶气！':
        cn[i-1] = cn[i-1] + ' ' + l
        cn.pop(i); break

kr = kr[kr.index('# 제1화'):]
assert len(cn) == len(kr), (len(cn), len(kr))

GREY = RGBColor(0x7A, 0x7A, 0x7A); BLACK = RGBColor(0x1A, 0x1A, 0x1A)
IDX = RGBColor(0xB0, 0xB0, 0xB0); ACC = RGBColor(0x8A, 0x5A, 0x10)
MD = re.compile(r'(\*\*[^*]+\*\*)')

NOTES = [
    '001-2 「卡尔马」 — 국명은 「卡马尔(카마르)」. 직전 회수본에서 지적한 표기 오류가 이 한 곳에 남아 있습니다.',
    '씬 헤더는 전부 「公主寝殿(공주 침전)」으로 통일됐는데, 지문·대사에는 「大殿(대전)」이 7회 남아 있습니다(003-2 Cut 이후·004-1). 원문 그대로 옮겼습니다.',
    '006-1 「脸色一僵吗」 — 「吗」는 「了」의 오타로 보입니다.',
]


def build(path, bilingual):
    d = docx.Document()
    sec = d.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(1.8)

    def para(space_after=0, space_before=0, keep=False, indent=0):
        p = d.add_paragraph(); pf = p.paragraph_format
        pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
        pf.line_spacing = 1.18; pf.keep_with_next = keep
        if indent: pf.left_indent = Pt(indent)
        return p

    def run(p, text, *, cnfont=False, size=11, color=BLACK, bold=False):
        r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold
        name = '微软雅黑' if cnfont else '맑은 고딕'
        r.font.name = name; r._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        return r

    def kr_line(p, text, size=11, color=BLACK):
        text = re.sub(r'\[\^\d\]', '', text)
        for seg in MD.split(text):
            if not seg: continue
            run(p, seg[2:-2] if seg.startswith('**') else seg,
                size=size, color=color, bold=seg.startswith('**'))

    # 표지
    t = para(space_after=2); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t, '거지 알라딘과 요술램프  ·  무료회차 1~7화', size=17, bold=True)
    t2 = para(space_after=4); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t2, '《乞丐阿拉丁与神灯》 第1~7集', cnfont=True, size=12, color=GREY)
    t3 = para(space_after=20); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = ('한국어 / 中文 원문 대조 합본  —  윗줄 = 한국어, 아랫줄 = 중국어 원문'
           if bilingual else '2026-08-05 작가 수정고 기준 한국어 번역본')
    run(t3, sub, size=9, color=GREY)

    scene_n = 0
    for k, c in zip(kr, cn):
        if k.startswith('# '):
            d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            h = para(space_after=10, keep=True)
            run(h, k[2:], size=16, bold=True)
            if bilingual: run(h, '   ' + c, cnfont=True, size=11, color=GREY)
            scene_n = 0
        elif k.startswith('### '):
            h = para(space_before=12, space_after=(2 if bilingual else 6), keep=True)
            run(h, k[4:], size=12.5, bold=True, color=ACC)
            if bilingual:
                h2 = para(space_after=6, keep=True)
                run(h2, c, cnfont=True, size=10, color=GREY)
            scene_n = 0
        elif k.startswith('인물'):
            p = para(space_after=(2 if bilingual else 8))
            kr_line(p, k, size=10, color=GREY)
            if bilingual:
                p2 = para(space_after=8); run(p2, c, cnfont=True, size=9, color=GREY)
        else:
            scene_n += 1
            if bilingual:
                p = para(space_after=1, keep=True)
                run(p, '%02d  ' % scene_n, size=8, color=IDX)
                kr_line(p, k)
                p2 = para(space_after=9, indent=20)
                run(p2, re.sub(r'^\*\*|\*\*', '', c), cnfont=True, size=9.5, color=GREY)
            else:
                p = para(space_after=(7 if k.startswith('△') else 5))
                kr_line(p, k)

    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h = para(space_after=8); run(h, '역주 (원문 표기)', size=13, bold=True, color=ACC)
    for n, t_ in enumerate(NOTES, 1):
        p = para(space_after=5); run(p, '%d. ' % n, size=10, color=GREY); kr_line(p, t_, size=10)

    d.save(path)
    print('OK', path)


build('거지 알라딘과 요술램프_무료회차 1-7화_한국어 번역본_v2.docx', False)
build('거지 알라딘과 요술램프_무료회차 1-7화_한중 대조본_v2.docx', True)
print('lines', len(kr))
