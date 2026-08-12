# -*- coding: utf-8 -*-
"""피칭 페이지 docx 빌더 — 거지 알라딘과 요술램프 (교정본 v2).

텍스트 원본 = 17_son_of_the_lamp_02_pitch_deck_v2.md (동일 내용).
"""
import os
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, '거지 알라딘과 요술램프_피칭페이지_v2.docx')

BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x7A, 0x7A, 0x7A)
ACC = RGBColor(0x8A, 0x5A, 0x10)
LINE = RGBColor(0xC9, 0xB8, 0x92)

d = docx.Document()
sec = d.sections[0]
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = sec.bottom_margin = Cm(1.9)


def para(space_after=0, space_before=0, keep=False, indent=0, spacing=1.30):
    p = d.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = spacing
    pf.keep_with_next = keep
    if indent:
        pf.left_indent = Pt(indent)
    return p


def run(p, text, *, size=11, color=BLACK, bold=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.font.name = '맑은 고딕'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return r


def section(title, en=None):
    """대제목 — 피칭 페이지 필드명."""
    p = para(space_before=17, space_after=7, keep=True)
    run(p, title, size=13, bold=True, color=ACC)
    if en:
        run(p, '  ' + en, size=9.5, color=GREY)


def sub(title):
    """소제목 — 섹션 안 논지 헤딩."""
    p = para(space_before=11, space_after=4, keep=True)
    run(p, title, size=11.5, bold=True)


def body(text, space_after=6, indent=0):
    p = para(space_after=space_after, indent=indent)
    run(p, text)


def field(label, value, gap=64):
    """라벨 + 값 한 줄."""
    p = para(space_after=5)
    p.paragraph_format.left_indent = Pt(gap)
    p.paragraph_format.first_line_indent = Pt(-gap)
    run(p, label, bold=True)
    run(p, '\t' + value)
    p.paragraph_format.tab_stops.add_tab_stop(Pt(gap))


def bullet(text, dim_tail=None):
    p = para(space_after=4, indent=11)
    p.paragraph_format.first_line_indent = Pt(-11)
    run(p, '· ', color=ACC, bold=True)
    run(p, text, size=10.5)
    if dim_tail:
        run(p, dim_tail, size=10.5, color=GREY)


def person(name, desc):
    p = para(space_before=9, space_after=3, keep=True)
    run(p, name, size=11, bold=True, color=ACC)
    body(desc, space_after=0)


def rule():
    p = para(space_before=6, space_after=2)
    run(p, '─' * 46, size=8, color=LINE)


# ── 표지 ─────────────────────────────────────────────
t = para(space_after=3)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t, '거지 알라딘과 요술램프', size=20, bold=True)

t2 = para(space_after=4)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t2, 'Beggar Aladdin and the Almighty Lamp', size=12, color=ACC)

t3 = para(space_after=6)
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t3, '피칭 페이지', size=9.5, color=GREY)

t4 = para(space_after=4)
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t4, '─' * 20, size=8, color=LINE)

# ── 타이틀 / 담당 ────────────────────────────────────
section('타이틀', 'Titles · タイトル')
field('영어', 'Beggar Aladdin and the Almighty Lamp')
field('한국어', '거지 알라딘과 요술램프')

section('담당 CM', '担当CM名')
body('Rowan Lee')

# ── 기본정보 ─────────────────────────────────────────
section('기본정보', 'Basic Information')
field('장르', 'Fantasy Action / Rom-com')
field('키워드', 'HIDDEN IDENTITY, OVERPOWERED, BODYGUARD')
field('회차', '50화 / 무료 EP1-7')
field('발화 언어', '영어')
field('제작 형식', 'AI 실사')
field('타겟층', '메인 = 북미·글로벌 남성향 18-45 / 서브 = 북미·글로벌 여성')
body('누구나 아는 알라딘 소재에 로맨스 축을 얹어, 남성향이면서 여성 시청까지 함께 흡수하는 '
     '중립 근접 남성향.', space_after=9, indent=64)

sub('로그라인')
body('무적의 삶에 질려 왕좌를 버린 진(Djinn)의 왕 알라딘. 왕권도 군단도 낡은 놋쇠 램프에 봉인해두고 '
     '인간 세상에서 백수로 구르다, 월급 200금화에 혹해 암살 표적이 된 공주의 전속 호위로 취직한다. '
     '온 왕국이 공주를 구한 정체불명의 고수를 찾는 동안, 그 고수는 공주 침전 기둥에 기대 사과를 씹고 있다.')

sub('레퍼런스')
bullet('StardustTV 《코드명: 블랙 호크》(61화·원제 Bukan Pengawal Biasa)',
       '  —  신분을 숨긴 최강자의 위장 경호 + 보호 로맨스 뼈대')
bullet('NetShort 《False Weakling, True Power》(50화)',
       '  —  힘 숨김 / 남성향 파워 판타지')
bullet('NetShort 《Swapped to a Beggar But He Is Apollo》(48화)',
       '  —  평민 위장 신급 남주 × 공주 로맨스 문법')
bullet('퍼블릭 도메인 《천일야화》 알라딘 설화',
       '  —  소재 인지도(원작료 0)')

# ── 피칭 사유 ────────────────────────────────────────
section('해당 IP 피칭 사유', 'Core Expectation')

p = para(space_after=8)
run(p, '왜 제작해야 하는지?   ·   시장성이 있는지?   ·   크리에이티브에 변별력이 있는지?',
    size=10, color=GREY)

rule()
p = para(space_before=4, space_after=6, keep=True)
run(p, '왜 제작해야 하는지?  시장성이 있는지?', size=11.5, bold=True, color=ACC)

sub('이집트 다음은 아랍 — 한 발 앞선 배경 선점')
body('DramaBox·NetShort에서 이집트 배경 작품이 이미 흥행한 케이스가 나왔다. 고대 오리엔트 비주얼이 '
     '버티컬 시장에서 먹힌다는 것은 검증된 상태다.')
body('여기서 한 발 앞서 아랍풍(사막 술탄국)을 시도한다. 이집트 흥행작들과 미술 자산의 성격이 가까워 '
     '제작 리스크는 낮으면서, 아직 경쟁작이 적은 구간이라 신선도는 확보된다.')
body('AI 숏폼 시장은 웨어울프 흥행이 끝나가는 국면이고, 새 판타지 소재와 배경의 유행은 빠르게 지나간다. '
     '지금 뜬 것을 그대로 따라가면 비글루가 출시할 즈음엔 이미 한물간 소재가 된다. 그래서 한 박자 먼저 간다.')

sub('설명이 필요 없는 소재')
body('애니메이션·실사 영화로 반복 영상화되며 인지도가 이미 최대치인 이야기다. 시청자에게 낯설지 않으니 '
     '진입 장벽이 낮다.')
body('남성향 작품이지만 진입 소재는 전 세계 누구나 아는 천일야화·알라딘이라, 남성향 판타지의 고질적 '
     "약점인 '세계관 학습 비용'이 처음부터 0이다. 램프를 문지르면 강력한 정령이 나온다는 규칙을 "
     '설명할 필요가 없어, 무료회차 시간을 전부 쾌감에 쓸 수 있다.')

sub('전 세계가 아는 이야기를, 판권료 0원으로')
body('알라딘·천일야화는 퍼블릭 도메인이라 원작료 지불이 없다. 저작권 이슈는 명칭·장면·의상 단위 '
     '금지 목록으로 회피한다(법률 검토 완료). 인지도는 전부 가져오고 권리 리스크만 지운 케이스다.')

rule()
p = para(space_before=4, space_after=6, keep=True)
run(p, '크리에이티브에 변별력이 있는지?', size=11.5, bold=True, color=ACC)

sub('알라딘과 램프의 역학 관계 뒤집기')
body('소원을 비는 알라딘이 아니라, 소원을 들어주는 쪽이 알라딘이다. 램프에 갇힌 쪽이 부하고, '
     '램프를 쥔 쪽이 알라딘 — 램프는 그가 자기 왕권과 군단을 통째로 봉인해 넣은 신물이다.')
body("'소원 세 개' 문법은 폐기했다. 그가 허리춤의 램프를 손끝으로 한 번 쓰다듬으면, 도망치던 자객이 "
     '통째로 램프 속으로 빨려 들어간다.')

sub('지니 설계')
body('정령 지니는 파란 개그 요정이 아니라 극존대 코미디 담당 시종장이다. 돌덩이 같은 흑빵을 왕궁 '
     '만찬처럼 차려 올리고, 주인이 접시 27장을 깨고 쫓겨난 것을 국란처럼 통탄한다.')
body('코미디를 정령이 전담해, 주인공은 끝까지 쿨하게 두고 톤만 가볍게 가져간다.')

sub('공주에게 기대하는 주체성은 그대로, 성깔을 얹는다')
body('나라와 백성을 걱정하는 주체성은 유지하되, 입이 먼저 나가는 성깔을 붙였다. 알라딘과 부딪힐 때 '
     '이 성깔이 그대로 로맨틱 코미디의 연료가 된다.')

sub('숏폼 히트작 문법의 조합')
body('한 히트작의 뼈대 위에, 다른 히트작에서 궁합이 맞는 부분만 골라 얹었다.')

# ── 인물 ─────────────────────────────────────────────
section('주요 인물 소개')

person('알라딘 (Aladdin, 남 / 외견 20대)  ·  남자주인공',
       '진 군단을 호령하던 정령들의 왕. 전쟁도 정복도 파괴도 순서까지 똑같이 반복되는 게 지겨워 '
       '"무적이야말로 가장 지루한 저주"라며 왕좌를 버렸다. 군단과 왕권을 통째로 낡은 놋쇠 램프에 '
       '봉인해 허리에 차고 인간 세상으로 내려왔지만, 세계 최강의 마법사도 생활력은 0이라 꼬치 장수에게 '
       '쫓겨나고 접시 27장을 깨고 잘린다. 월급 200금화짜리 공주 호위 공고를 보고 "공주를 지키는 건 '
       '양심 있는 자의 의무"라며 냉큼 지원한다. 시키기 전엔 절대 나서지 않고, 시키면 뺨 한 대로 끝낸다.')

person('소라야 (Soraya, 여)  ·  여자주인공',
       '카마르(Qamar) 왕국의 공주. 술탄이 쓰러진 뒤 외척 총독 마리크의 왕권 찬탈에 홀로 맞서고 있다. '
       '"내가 살아 있는 한 카마르는 썩은 피가 넘볼 수 없다"고 면전에서 쏘아붙일 만큼 강하지만, '
       '마리크가 보낸 폭도들이 먹인 약에 취해 끌려가던 밤만은 평생의 수치로 남았다. 그 밤 자기를 구한 '
       '남자를 다시 만났을 때, 하필 그가 자기 침전에 배치된 한심한 호위라는 게 문제다. 자객이 사라질 '
       '때마다 그를 의심하지만, 사과나 씹고 있는 얼굴을 보면 매번 결론이 뒤집힌다.')

person('지니 (Genie, 남)  ·  램프의 정령',
       '램프에 봉인된 진의 왕 직속 시종장. 주인이 곤란해질 때마다 램프에서 금빛 연기와 함께 제멋대로 '
       '피어올라 주인을 더 곤란하게 만든다. 왕궁 의전이 몸에 밴 극진한 존대로 거지꼴 주인을 모시는 게 '
       '그의 비극이자 이 극의 개그. "인간 세계에선 그 호칭 금지, 끄덕이면 예스 흔들면 노"라는 규칙 '
       '때문에 매번 말이 막힌다. 궁의 이상 징후를 먼저 감지해 경고하지만 결정타는 언제나 왕의 손에 남긴다.')

person('마리크 (Malik, 남)  ·  메인 빌런',
       '카마르 왕국의 외척 총독. 술탄이 혼수상태에 빠진 틈을 타 국정을 쥐고, 귀족 할리드와의 정략혼으로 '
       '공주를 치우려 한다. 정면으로 안 되면 그림자 자객을, 자객이 지워지면 고위 마법사를, 그마저 '
       '무너지면 여러 왕국이 현상금을 건 찬탈자를 불러들인다. 판이 어그러질 때마다 "술탄이 숨겨둔 '
       '대마법사가 있다"고 확신하지만, 정작 매일 공주 옆에 서 있는 호위는 눈에 들어오지도 않는다.')

person('할리드 (Khalid, 남)  ·  귀족 공자 / 연적',
       '마리크가 미는 공주의 약혼자. 보석을 주렁주렁 달고 다니며 고위 마법사 둘을 개인 호위로 끌고 '
       '다닌다. 공주의 침전 문을 부수고 들어와 "내 손에서 한 수만 버티면 돌아가주지"라고 지껄이다가, '
       '이름도 모르는 호위에게 뺨 한 대를 맞고 속바지 차림으로 문밖으로 날아간다. 굴욕을 안고 마리크에게 '
       '달려가 더 큰 것을 불러온다.')

person('바스마 (Basma, 여)  ·  공주의 시녀',
       '호위 선발을 담당한 공주의 최측근. 전 종목 0점을 받은 알라딘을 "그래도 공주님이 악몽은 안 '
       '꾸시겠네"라며 얼굴 하나로 뽑아버린 장본인. 사태가 터질 때마다 몸으로 공주를 막아서는 인물이자, '
       '이 극의 리액션 담당.')

person('레이븐 (Raven, 남)  ·  찬탈자',
       '여러 왕국이 현상금을 건 지명수배자. 창백한 얼굴 전체에 뒤틀린 검은 룬이 새겨져 있다. 마리크가 '
       '마지막 카드로 불러들인 존재이며, 7화 마지막 컷에 등장해 다음 상대를 예고한다.')

# ── 줄거리 ───────────────────────────────────────────
section('줄거리')
body('진 군단을 호령하던 왕 알라딘은 무적의 삶에 질려 왕좌를 버린다. 왕권과 군단을 낡은 놋쇠 램프에 '
     '봉인하고 내려온 인간 세상에서 원하는 건 월급과 조용한 삶뿐. 하지만 세계 최강의 마법사도 생활력은 '
     '0이라 접시나 깨고 쫓겨나던 어느 밤, 뒷골목에서 폭도에게 끌려가던 카마르의 공주 소라야를 구한다. '
     "다음 날 그는 '월급 200금화'에 혹해 하필 그 공주의 전속 호위로 취직한다. 궁은 이미 전쟁터다. "
     '술탄이 쓰러진 틈을 타 외척 총독 마리크가 정략혼과 자객으로 왕권을 삼키려 하고, 알라딘은 날아드는 '
     '단검을 손가락 하나로 꺾고 자객을 램프에 삼키고 귀족 할리드의 고위 마법사들을 뺨 한 대씩으로 '
     "눕히며 판을 뒤집는다. 왕국은 '술탄이 숨겨둔 대마법사'를 찾아 술렁이지만, 그가 진의 왕이라는 것은 "
     '화면 밖 관객만 안다. 마리크가 찬탈자 레이븐을 끌어들여 왕궁을 통째로 엎으려 하고, 소라야는 그날 '
     '밤 자기를 구한 손이 매일 옆에서 사과나 씹던 손이라는 것을 알게 된다. 만인 앞에 정체를 드러낸 '
     '알라딘이 마리크와 레이븐을 무너뜨리고 술탄이 지켜보는 앞에서 공주의 손을 잡는 날 — 사막 저편, '
     '비워둔 진의 왕좌가 그를 부른다.')

# ── 트리트먼트 ───────────────────────────────────────
section('초반 회차 트리트먼트', '1화~무료 마지막 화')
p = para(space_after=9)
run(p, '무료회차 대본 1-7화 첨부', size=10, color=GREY)

TREATMENT = [
    ('1화',
     '진의 왕 알라딘이 손가락 하나로 적 대군을 통째로 지운다. "무적이란 건 가장 지루한 저주였군." '
     '군단을 램프에 봉인하고 내려간 인간 세상, 그날 밤 폐가에서 돌덩이 같은 흑빵을 씹는다.',
     '골목에서 비명. 약을 먹인 공주를 마차로 끌고 가려던 폭도들 앞에 반쪽 흑빵을 든 거지 청년이 선다.'),
    ('2화',
     '금빛 몇 줄기에 폭도 전원이 벽에 처박힌다. 다음 날 아침 마차 안, 얼굴이 붉어진 소라야가 금화 '
     '주머니를 던진다. "어젯밤 일은 없었던 걸로 해. 한마디라도 흘리면 죽인다."',
     '광장 게시판 [공주 전속 호위 고액 모집]. "안 가!" 돌아서던 알라딘이 "월급이 200금화입니다" '
     '한마디에 게시물을 뜯어낸다. "공주를 지키는 건 양심 있는 자의 의무지."'),
    ('3화',
     '전 종목 0점을 받은 알라딘이 얼굴 하나로 뽑힌다. 침전에서 재회한 소라야 앞에 허리를 90도로 '
     '꺾는다. "사람을 잘못 보셨습니다. 여신 같은 여자는 본 적도 없습니다."',
     '"전하, 이 방에 쥐가 들었습니다." 구석의 그림자가 찢어지고 단검이 소라야의 등으로 날아간다.'),
    ('4화',
     '손가락 한 번에 단검 궤도가 꺾이고, 도망치던 자객은 램프에 통째로 삼켜진다. "방금… 네가 한 '
     '거야?" "제가요? 저 방금 오줌 쌀 뻔했는데요."',
     '마리크의 밀실, 푸른 촛불이 꺼진다. "암영을 지웠다면 최소 대마법사급… 이제 할리드 그 멍청이를 '
     '써야겠어."'),
    ('5화',
     '침전 문이 부서지며 할리드가 고위 마법사 둘을 끌고 들어온다. "네 호위 중에 내 손에서 한 수라도 '
     '버티는 놈이 있으면 돌아가주지."',
     '사과 씹는 소리가 침전에 크게 울린다. 기둥에 기대 구경하던 호위를 발견한 소라야가 손가락으로 '
     '그를 가리킨다. "너! 알라딘, 이 개자식! 가서 저것들 다 죽여버려!"'),
    ('6화',
     '"저요?" "그래, 너!" "절차가 맞았군요. 업무 시작." 알라딘이 공주 앞을 막아선다. "쫓아내면 '
     '됩니까, 기절시킬까요, 아니면 앞으로 문만 봐도 무서워지게 만들까요?"',
     '두 마법사가 동시에 주문을 시작하고 발밑에서 검은 법진이 타오른다.'),
    ('7화',
     '불뱀이 침전을 채우는 순간 알라딘이 사라진다. 짝—! 뺨 네 번에 마법사와 호위 전원이 벽에 '
     '처박히고, 허리띠가 끊긴 할리드가 붉은 속바지 차림으로 문밖으로 날아간다. "꺼져."',
     '마리크의 저택. 뺨이 부어터진 할리드 앞으로 검은 룬을 새긴 찬탈자 레이븐이 걸어 들어온다. '
     '"안녕들 하신가. 누가 자잘한 장애물을 좀 치워달라던데."'),
]

for ep, line, cliff in TREATMENT:
    h = para(space_before=10, space_after=3, keep=True)
    run(h, ep, size=11, bold=True, color=ACC)
    body(line, space_after=3)
    p = para(space_after=0)
    run(p, '클리프  ', size=10.5, bold=True, color=GREY)
    run(p, cliff, size=10.5)

d.save(OUT)
print('OK', OUT)
